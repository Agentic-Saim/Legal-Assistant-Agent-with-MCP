"""
PDF and document parsing utilities for legal document processing.
Supports PDF, DOCX, and plain text formats.
"""
import io
import re
from typing import Dict, List, Optional, Any
from pathlib import Path
from dataclasses import dataclass


@dataclass
class ParsedDocument:
    """Structure for parsed document content."""
    full_text: str
    sections: Dict[str, str]
    party_names: List[str]
    effective_date: Optional[str]
    signature_blocks: List[str]
    defined_terms: Dict[str, str]
    clause_numbers: Dict[str, str]
    metadata: Dict[str, Any]


class DocumentParser:
    """
    Multi-format document parser for legal documents.
    Supports PDF, DOCX, and plain text.
    """
    
    # Common legal section patterns
    SECTION_PATTERNS = {
        "parties": r"(?:parties|recitals|background)\s*[:\-]?",
        "definitions": r"(?:definitions|defined terms|interpretation)\s*[:\-]?",
        "scope_of_work": r"(?:scope of work|services|statement of work)\s*[:\-]?",
        "payment": r"(?:payment|compensation|fees|consideration)\s*[:\-]?",
        "intellectual_property": r"(?:intellectual property|IP|ownership|work product)\s*[:\-]?",
        "confidentiality": r"(?:confidentiality|non-disclosure|confidential information)\s*[:\-]?",
        "non_compete": r"(?:non-compete|non-solicitation|restrictive covenant)\s*[:\-]?",
        "representations": r"(?:representations|warranties|reps and warranties)\s*[:\-]?",
        "indemnification": r"(?:indemnification|indemnity|hold harmless)\s*[:\-]?",
        "liability": r"(?:limitation of liability|liability cap|damages)\s*[:\-]?",
        "termination": r"(?:termination|expiration|expiry)\s*[:\-]?",
        "dispute_resolution": r"(?:dispute resolution|governing law|jurisdiction|arbitration)\s*[:\-]?",
        "force_majeure": r"(?:force majeure|act of god|impossibility)\s*[:\-]?",
        "miscellaneous": r"(?:miscellaneous|boilerplate|general provisions)\s*[:\-]?",
    }
    
    # Party name patterns
    PARTY_PATTERNS = [
        r'"([^"]+)"(?:,?\s+(?:a\s+)?(?:corporation|LLC|LLP|Inc\.|Ltd\.|Company))',
        r'(?:hereinafter|referred to as)\s+"([^"]+)"',
        r'^([A-Z][A-Za-z\s&\.]+),?\s+(?:a\s+)?(?:corporation|LLC|LLP|Inc\.|Ltd\.)',
    ]
    
    # Date patterns
    DATE_PATTERNS = [
        r'(?:effective|dated|dated as of|as of)\s+([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})',
        r'(\d{1,2}/\d{1,2}/\d{4})',
        r'(\d{4}-\d{2}-\d{2})',
    ]
    
    # Defined term pattern
    DEFINED_TERM_PATTERN = r'"([^"]+)"\s+(?:means|shall mean|refers to|has the meaning set forth below)'
    
    # Clause number pattern
    CLAUSE_NUMBER_PATTERN = r'^(?:Section|Article|Clause|§)\s*([\d\.]+[A-Za-z]?)'
    
    def __init__(self):
        self.pdfplumber = None
        self.pymupdf = None
        self.docx = None
        self._libraries_checked = False
    
    def _check_libraries(self):
        """Lazy import of document processing libraries."""
        if self._libraries_checked:
            return
        
        try:
            import pdfplumber
            self.pdfplumber = pdfplumber
        except ImportError:
            pass
        
        try:
            import fitz as pymupdf
            self.pymupdf = pymupdf
        except ImportError:
            pass
        
        try:
            from docx import Document
            self.docx = Document
        except ImportError:
            pass
        
        self._libraries_checked = True
    
    def parse_file(self, file_path: str) -> ParsedDocument:
        """
        Parse a document file (PDF, DOCX, or TXT).
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParsedDocument with extracted content
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        self._check_libraries()
        
        if path.suffix.lower() == '.pdf':
            return self._parse_pdf(path)
        elif path.suffix.lower() in ['.docx', '.doc']:
            return self._parse_docx(path)
        elif path.suffix.lower() == '.txt':
            return self._parse_text(path)
        else:
            # Try to parse as text
            return self._parse_text(path)
    
    def parse_text(self, text: str, document_name: str = "Document") -> ParsedDocument:
        """
        Parse raw text content.
        
        Args:
            text: Raw document text
            document_name: Name/identifier for the document
            
        Returns:
            ParsedDocument with extracted content
        """
        sections = self._extract_sections(text)
        parties = self._extract_parties(text)
        effective_date = self._extract_effective_date(text)
        signature_blocks = self._extract_signature_blocks(text)
        defined_terms = self._extract_defined_terms(text)
        clause_numbers = self._extract_clause_numbers(text)
        
        return ParsedDocument(
            full_text=text,
            sections=sections,
            party_names=parties,
            effective_date=effective_date,
            signature_blocks=signature_blocks,
            defined_terms=defined_terms,
            clause_numbers=clause_numbers,
            metadata={
                "document_name": document_name,
                "word_count": len(text.split()),
                "character_count": len(text),
            }
        )
    
    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """Parse PDF file using pdfplumber or PyMuPDF."""
        text = ""
        
        # Try pdfplumber first (better text extraction)
        if self.pdfplumber:
            try:
                with self.pdfplumber.open(str(path)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                # Fall back to PyMuPDF
                pass
        
        # Try PyMuPDF if pdfplumber failed or isn't available
        if not text and self.pymupdf:
            try:
                doc = self.pymupdf.open(str(path))
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except Exception as e:
                pass
        
        if not text:
            raise ValueError("Could not extract text from PDF. File may be image-only.")
        
        return self.parse_text(text, document_name=path.name)
    
    def _parse_docx(self, path: Path) -> ParsedDocument:
        """Parse DOCX file using python-docx."""
        if not self.docx:
            raise ImportError("python-docx not installed. Cannot parse DOCX files.")
        
        doc = self.docx(path)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\t"
                text += "\n"
        
        return self.parse_text(text, document_name=path.name)
    
    def _parse_text(self, path: Path) -> ParsedDocument:
        """Parse plain text file."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # Try with latin-1 encoding
            with open(path, 'r', encoding='latin-1') as f:
                text = f.read()
        
        return self.parse_text(text, document_name=path.name)
    
    def _extract_sections(self, text: str) -> Dict[str, str]:
        """Extract document sections based on legal section patterns."""
        sections = {}
        text_lower = text.lower()
        
        # Find all section headers
        section_positions = []
        for section_name, pattern in self.SECTION_PATTERNS.items():
            matches = list(re.finditer(pattern, text_lower, re.IGNORECASE))
            for match in matches:
                section_positions.append((match.start(), section_name, match.group()))
        
        # Sort by position
        section_positions.sort(key=lambda x: x[0])
        
        # Extract section content
        for i, (pos, section_name, header) in enumerate(section_positions):
            # Find end position (start of next section or end of document)
            if i + 1 < len(section_positions):
                end_pos = section_positions[i + 1][0]
            else:
                end_pos = len(text)
            
            # Extract section text
            section_text = text[pos:end_pos].strip()
            
            # Store if not already present (prefer first occurrence)
            if section_name not in sections:
                sections[section_name] = section_text
        
        return sections
    
    def _extract_parties(self, text: str) -> List[str]:
        """Extract party names from document."""
        parties = []
        
        for pattern in self.PARTY_PATTERNS:
            matches = re.findall(pattern, text, re.MULTILINE | re.IGNORECASE)
            parties.extend(matches)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_parties = []
        for party in parties:
            if party not in seen and len(party) > 2:
                seen.add(party)
                unique_parties.append(party)
        
        return unique_parties[:10]  # Limit to 10 parties
    
    def _extract_effective_date(self, text: str) -> Optional[str]:
        """Extract effective date from document."""
        for pattern in self.DATE_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1)
        return None
    
    def _extract_signature_blocks(self, text: str) -> List[str]:
        """Extract signature blocks from document."""
        signature_blocks = []
        
        # Look for signature patterns
        patterns = [
            r"(?:IN WITNESS WHEREOF|EXECUTED|SIGNED|AGREED AND ACCEPTED)[\s\S]{0,500}(?:By:|Signature:|Signed:)[\s\S]{0,200}",
            r"(?:By:|Signature:)\s*\n\s*[A-Za-z\s]+\s*\n\s*(?:Title:|Name:|Date:)[\s\S]{0,100}",
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            signature_blocks.extend(matches)
        
        return signature_blocks[:5]  # Limit to 5 signature blocks
    
    def _extract_defined_terms(self, text: str) -> Dict[str, str]:
        """Extract defined terms and their definitions."""
        defined_terms = {}
        
        # Find "Term" means definition pattern
        pattern = r'"([^"]+)"\s+(?:means|shall mean|refers to)\s+([^\.]+)\.'
        matches = re.findall(pattern, text, re.IGNORECASE)
        
        for term, definition in matches:
            defined_terms[term] = definition.strip()
        
        # Also look for section-based definitions
        definitions_section = None
        for section_name, content in self._extract_sections(text).items():
            if 'definition' in section_name:
                definitions_section = content
                break
        
        if definitions_section:
            # Look for capitalized term followed by definition
            term_pattern = r'"([A-Z][A-Za-z]+)"\s+(?:means|shall mean|refers to)\s+([^\.]+)'
            matches = re.findall(term_pattern, definitions_section)
            for term, definition in matches:
                if term not in defined_terms:
                    defined_terms[term] = definition.strip()
        
        return defined_terms
    
    def _extract_clause_numbers(self, text: str) -> Dict[str, str]:
        """Extract clause/section numbers and their content."""
        clause_numbers = {}
        
        lines = text.split('\n')
        current_clause = None
        current_content = []
        
        for line in lines:
            match = re.match(self.CLAUSE_NUMBER_PATTERN, line.strip(), re.IGNORECASE)
            if match:
                # Save previous clause
                if current_clause:
                    clause_numbers[current_clause] = '\n'.join(current_content)
                
                # Start new clause
                current_clause = match.group(1)
                current_content = [line]
            elif current_clause:
                current_content.append(line)
        
        # Save last clause
        if current_clause:
            clause_numbers[current_clause] = '\n'.join(current_content)
        
        return clause_numbers


# Singleton instance
_parser = DocumentParser()


def parse_document(file_path: str) -> ParsedDocument:
    """
    Parse a document file.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        ParsedDocument with extracted content
    """
    return _parser.parse_file(file_path)


def parse_text_content(text: str, document_name: str = "Document") -> ParsedDocument:
    """
    Parse raw text content.
    
    Args:
        text: Raw document text
        document_name: Name/identifier for the document
        
    Returns:
        ParsedDocument with extracted content
    """
    return _parser.parse_text(text, document_name)
